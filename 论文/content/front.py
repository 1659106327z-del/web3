"""封面 / 中文摘要 / Abstract / 目录 (占位)"""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def write(doc, api):
    _cover(doc, api)
    _zh_abstract(doc, api)
    _en_abstract(doc, api)
    _toc(doc, api)


def _cover(doc, api):
    """封面：题目居中 + 学生信息"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("江西农业大学")
    run.font.name = "黑体"
    run.font.size = Pt(22)
    run.bold = True
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    from docx.oxml import OxmlElement
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本科毕业论文（设计）")
    run.font.size = Pt(20)
    run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Web 的操作系统进程调度算法可视化演示系统的设计与开发")
    run.font.size = Pt(20)
    run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Design and Development of a Web-Based Visualization System for Operating System Process Scheduling Algorithms")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    for _ in range(8):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("")

    info_lines = [
        ("学生姓名：", "（姓名）"),
        ("学　　号：", "6020225470"),
        ("专业班级：", "计算机 2201"),
        ("指导教师：", "周双娥"),
        ("提交日期：", "2026 年 5 月"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + value)
        run.font.size = Pt(14)
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rpr.append(rfonts)


def _zh_abstract(doc, api):
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("摘  要")
    run.font.size = Pt(14)
    run.bold = True
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    api["para"](
        doc,
        "操作系统进程调度是计算机专业本科教学的核心内容，但抽象的算法逻辑长期依赖静态板书与 PPT 配图，"
        "学生难以直观感受调度过程中进程在多个队列之间的迁移、上下文切换以及指标变化。"
        "针对这一痛点，本研究基于 Web 全栈技术构建了一套可视化演示系统，希望以"
        "「动手即可观察」的方式降低算法学习门槛。"
    )
    api["para"](
        doc,
        "系统采用 Next.js 14 App Router、React 18 与 TypeScript 作为前端基底，配合 Tailwind CSS 完成了暖色调"
        "玻璃拟态（Glassmorphism）的视觉风格；状态层使用 Zustand 维护播放控制与仿真上下文；"
        "可视化层将动态甘特图（HTML5 Canvas）与多区调度舞台（SVG + Framer Motion）相结合，"
        "实现了从调度事件序列到画面位置的派生式渲染。"
        "核心仿真引擎覆盖先来先服务、短作业优先、最短剩余时间优先、时间片轮转、优先级调度与多级反馈队列六种经典算法，"
        "以统一的事件驱动契约暴露给视图层。"
        "针对教学与课程实验的衍生需求，系统额外实现了 MFQ 设计器、双算法/六算法矩阵对比与五维雷达图等三个扩展模块，"
        "并通过 Prisma 6 与 SQLite 配合自实现的 jose JWT 鉴权完成跨设备账号同步与个人用例的持久化。"
        "项目已经通过 Railway 完成容器化部署，可在任意联网设备上访问。"
    )
    api["para"](
        doc,
        "测试环节采用功能、兼容性与性能三类用例。功能层面使用真实的 HTTP 请求验证了登录、注册、用例 CRUD 等关键链路；"
        "兼容性层面在 375、768、1440 三档断点下逐项检查了响应式表现；"
        "性能层面以一组 5 进程的标准用例对六种算法的平均周转时间、平均等待时间、CPU 利用率与上下文切换次数做了量化对比，"
        "并以 MFQ 设计器分别套用等差递增、等比递增与默认配置三种参数，"
        "验证了多维评分系统对调度配置敏感度的区分能力。"
        "结果表明系统在交互流畅度与教学表达力两方面均达到了预期目标。"
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("关键词：")
    run.font.size = Pt(14)
    run.bold = True
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    run = p.add_run("操作系统；进程调度；算法可视化；Web 全栈；玻璃拟态")
    run.font.size = Pt(12)
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rpr.append(rfonts)


def _en_abstract(doc, api):
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = "Times New Roman"

    paragraphs = [
        "Process scheduling is among the most algorithmically dense topics in operating system courses, "
        "yet undergraduate teaching still relies heavily on chalkboard sketches and static presentation slides. "
        "Students often struggle to picture how processes migrate between queues, how context switches accumulate, "
        "and how metrics evolve over time. This thesis presents a web-based visualization system designed to "
        "lower that barrier by letting learners watch and tweak the scheduler in real time.",
        "The front end is built on Next.js 14 with the App Router, React 18 and TypeScript, paired with "
        "Tailwind CSS to deliver a warm-toned glassmorphism interface. Playback and simulation context are "
        "kept in Zustand stores, while the rendering layer combines a Canvas-based dynamic Gantt chart with "
        "an SVG stage animated by Framer Motion. The simulation engine implements six classical algorithms "
        "(FCFS, SJF, SRTF, RR, PSA, MFQ) behind a uniform event-driven contract, so the same data feeds both "
        "the visualization and the metrics dashboard.",
        "Beyond the standard suite the system ships three teaching-oriented extensions: an MFQ designer "
        "where users tailor the number of queues and the per-level quanta and receive a five-dimensional "
        "score against the textbook baseline; a comparison page that lays six algorithms side by side and "
        "renders a radar plot of normalized metrics; and a cross-device account layer powered by Prisma 6, "
        "SQLite, bcrypt and a custom HS256 JWT issued by edge middleware. The project is deployed on "
        "Railway and can be reached from any networked device.",
        "Testing covered functional, responsive and performance dimensions. Authentication, preset CRUD "
        "and other key paths were exercised through real HTTP traffic; the interface was inspected at 375, "
        "768 and 1440 pixel breakpoints; six algorithms were profiled on a shared five-process workload to "
        "compare average turnaround time, waiting time, CPU utilisation and context switches. A further "
        "experiment ran arithmetic, geometric and default quanta through the MFQ designer and confirmed "
        "the scoring system reflects scheduling intent. Overall the system meets the educational and "
        "interactive goals set out at the project proposal stage.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = Pt(20)
        from docx.enum.text import WD_LINE_SPACING
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run("Key words: ")
    run.font.size = Pt(14)
    run.bold = True
    run.font.name = "Times New Roman"
    run = p.add_run("Operating System; Process Scheduling; Algorithm Visualization; Full-stack Web; Glassmorphism")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"


def _toc(doc, api):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    run.font.size = Pt(14)
    run.bold = True
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("（请在 Microsoft Word 中将光标置于此处，点击「引用 → 目录 → 自动目录 1」即可生成；脚本不再人工维护目录条目，避免与正文实际页码不一致。）")
    run.font.size = Pt(11)
    run.italic = True
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rpr.append(rfonts)
