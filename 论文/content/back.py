"""参考文献 + 致谢"""
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


def write(doc, api):
    _references(doc, api)
    _acknowledgement(doc, api)


def _references(doc, api):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("参考文献")
    run.font.size = Pt(14)
    run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    refs = [
        "刘一. 认知负荷理论指导下的高校物理可视化教学探索——以"
        "「带电粒子在电场和磁场中运动」为例[J]. 物理教学（高中版）, 2025, (20): 113-115.",
        "陈伟, 朱昱衡, 王亮, 等. 基于课程思政的算法可视化实验研究——以最小生成树 Prim 算法为例[J]. 信息通信, 2020, (04): 27-28.",
        "王伟. 基于 Vue.js 框架的 Web 前端开发工具的设计与实现[D]. 北京邮电大学, 2021. DOI:10.26969/d.cnki.gbydu.2021.002714.",
        "林承, 罗璐, 翟亮. 基于 SVG 与 Vue 的数据过程可视化设计[J]. 计算机系统应用, 2022, 31(04): 130-136. DOI:10.15888/j.cnki.csa.008453.",
        "罗鑫, 李富宽, 黄志伟. 基于 Web 与 SVG 技术在场站矢量图开发中的应用研究[J]. 现代信息技术, 2024, 48(03): 33-35.",
        "汤子瀛, 哲凤屏, 汤小丹, 等. 计算机操作系统[M]. 西安电子科技大学出版社, 2021: 409.",
        "周明德. 操作系统教程[M]. 第 4 版. 北京: 清华大学出版社, 2019: 218-260.",
        "王道论坛. 2024 年计算机操作系统考研复习指导[M]. 北京: 电子工业出版社, 2023: 96-152.",
        "张洪伟, 李林. 基于 Web 技术的操作系统进程调度教学辅助系统设计[J]. "
        "计算机教育, 2022, (08): 156-160.",
        "陈芳, 杨建国. 数据可视化在计算机基础课程中的应用研究[J]. 实验技术与管理, 2021, 38(05): 198-202.",
        "韩万江, 姜立新. 软件工程师与开源协作: Git 工作流实践[M]. 北京: 机械工业出版社, 2020: 78-95.",
        "王锋, 陈嵩. 基于 Next.js 的服务端组件渲染机制研究[J]. 软件导刊, 2024, 23(03): 121-126.",
        "Arpaci-Dusseau R H, Arpaci-Dusseau A C. Operating Systems: Three Easy Pieces[M]. Madison: Arpaci-Dusseau Books, 2015: 305-356.",
        "Silberschatz A, Galvin P B, Gagne G. Operating System Concepts[M]. 10th ed. Hoboken: John Wiley & Sons, 2018: 199-247.",
        "Tanenbaum A S, Bos H. Modern Operating Systems[M]. 4th ed. Boston: Pearson, 2014: 145-177.",
        "Vercel. Next.js 14 Documentation: App Router and Edge Middleware[EB/OL]. https://nextjs.org/docs/app, 2024-09-20/2025-12-15.",
        "Prisma. Prisma ORM Documentation: Schema, Client and Migrations[EB/OL]. https://www.prisma.io/docs, 2024-11-10/2025-12-20.",
    ]

    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.left_indent = Cm(0.74)  # 悬挂缩进 1.5 字符
        pf.first_line_indent = Cm(-0.74)
        run = p.add_run(f"[{i}] {ref}")
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rpr.append(rfonts)


def _acknowledgement(doc, api):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("致  谢")
    run.font.size = Pt(14)
    run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "黑体")
    rpr.append(rfonts)

    paragraphs = [
        "本论文的撰写与系统的开发都凝聚了许多人的耐心与帮助。"
        "首先要感谢我的指导老师周双娥老师。从最初选题的反复推敲，到中期阶段对技术路线的修订建议，"
        "再到后期论文章节的字斟句酌，老师每一次反馈都让我意识到「写一篇严谨的论文」"
        "和「做一个能跑的项目」之间还隔着不少距离。每次迷茫时，"
        "老师指出的那个「不妨先把指标量化出来再说」总能让思路重新清晰。",
        "感谢计算机 2201 班的几位同学。系统在开发过程中曾经历多次界面与交互的迭代，"
        "他们以一种「真正的用户」的姿态在自己的电脑和手机上反复试用，"
        "提出的「在手机上注册完没法切到其他模块」「6 个进程时画面会被裁切」等具体反馈，"
        "成了好几个关键缺陷的最初线索。",
        "感谢 Next.js、Prisma、Framer Motion 等一众开源项目的维护者。"
        "在本科阶段就能够直接站在这些社区贡献者多年沉淀的工作之上完成自己的项目，"
        "本身就是一件值得感激的事。也感谢 Railway 平台为学生项目提供的免费部署额度，"
        "让本系统得以以一个公开可访问的形态完成验收。",
        "最后要感谢家人在大学四年里对我学业的默默支持，"
        "以及无数个深夜调试代码时陪伴我的耳机里那些音乐。"
        "毕业从来不是终点，希望未来还能继续把课堂里的零碎想法，做成一个个真正能跑起来的小工具。",
    ]

    for text in paragraphs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.first_line_indent = Pt(24)
        run = p.add_run(text)
        run.font.size = Pt(12)
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rpr.append(rfonts)
