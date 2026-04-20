"""
毕业论文 docx 生成主脚本

按江西农业大学本科毕业论文（设计）撰写要求（附件 1/2）排版。
依赖：pip install python-docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# 把 content/ 加到 import 路径
sys.path.insert(0, str(Path(__file__).parent))

from content import front, ch1_intro, ch2_tech, ch3_req, ch4_design, ch5_impl, ch6_test, ch7_summary, back  # noqa: E402

OUTPUT = Path(__file__).parent / "基于Web的操作系统进程调度算法可视化演示系统的设计与开发.docx"


# ---------- 字体与样式辅助 ----------

def set_run_font(run, *, font_zh="宋体", font_en="Times New Roman", size_pt: float = 12, bold: bool = False, color: tuple[int, int, int] | None = None):
    """中文与西文字体分离设置；size_pt 用 docx 标号约定（小四 = 12，四号 = 14，五号 = 10.5）"""
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_en)
    rfonts.set(qn("w:hAnsi"), font_en)
    rfonts.set(qn("w:eastAsia"), font_zh)


def set_paragraph_format(p, *, line_spacing_pt: float = 20, first_line_indent_chars: float = 0, alignment=None, space_before_lines: float = 0, space_after_lines: float = 0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent_chars > 0:
        # 首行缩进按字符（小四号约 12pt，每字符 ≈ 12pt 宽度）
        pf.first_line_indent = Pt(12 * first_line_indent_chars)
    if space_before_lines > 0:
        pf.space_before = Pt(20 * space_before_lines)
    if space_after_lines > 0:
        pf.space_after = Pt(20 * space_after_lines)


# ---------- 公开 API：注入到 content 模块使用 ----------

def add_h1(doc, text: str, *, page_break: bool = True):
    """一级标题：黑体四号，左顶格，前后段间距 1 行"""
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before_lines=1, space_after_lines=1)
    run = p.add_run(text)
    set_run_font(run, font_zh="黑体", font_en="Times New Roman", size_pt=14)
    return p


def add_h2(doc, text: str):
    """二级标题：黑体小四号"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before_lines=0.5, space_after_lines=0.3)
    run = p.add_run(text)
    set_run_font(run, font_zh="黑体", font_en="Times New Roman", size_pt=12)
    return p


def add_h3(doc, text: str):
    """三级标题：宋体小四号"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before_lines=0.3, space_after_lines=0.2)
    run = p.add_run(text)
    set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=12, bold=True)
    return p


def add_para(doc, text: str, *, indent: bool = True):
    """正文：宋体小四号，行距固定值 20 磅，首行缩进 2 字符"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, first_line_indent_chars=2 if indent else 0)
    run = p.add_run(text)
    set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=12)
    return p


def add_code_block(doc, code: str):
    """代码块：等宽字体，宋体五号样式"""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing_pt=15, first_line_indent_chars=0)
        run = p.add_run(line)
        set_run_font(run, font_zh="Consolas", font_en="Consolas", size_pt=10)


def add_caption(doc, text: str, *, kind: str = "图"):
    """图题或表题：宋体五号居中"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
    return p


def add_figure_placeholder(doc, caption: str):
    """图位占位：方框 + 图题"""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("[ 图位 · 实际部署后从 web 截图替换 ]")
    set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=10.5)
    add_caption(doc, caption, kind="图")


def add_table(doc, headers: list[str], rows: list[list[str]], caption: str | None = None):
    """三线表：仅顶部、底部为粗线，标题行下为细线"""
    if caption:
        # 表题置于表上方
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing_pt=20, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(caption)
        set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # 写表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, line_spacing_pt=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(h)
        set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)

    # 写数据
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, line_spacing_pt=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(str(val))
            set_run_font(run, font_zh="宋体", font_en="Times New Roman", size_pt=10.5)

    # 三线表边框：清掉默认，再加顶/底/标题下三条
    tbl_pr = table._element.tblPr
    for borders in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "bottom"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")  # 粗线
        b.set(qn("w:color"), "000000")
        borders.append(b)
    for name in ("left", "right", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        borders.append(b)
    inside_h = OxmlElement("w:insideH")
    inside_h.set(qn("w:val"), "single")
    inside_h.set(qn("w:sz"), "4")  # 细线
    inside_h.set(qn("w:color"), "000000")
    borders.append(inside_h)
    tbl_pr.append(borders)

    return table


# ---------- 页面与页码 ----------

def setup_page(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._element.append(fld)
    set_run_font(run, font_zh="Times New Roman", font_en="Times New Roman", size_pt=10.5)


# ---------- 默认样式调整 ----------

def normalize_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def main():
    doc = Document()
    normalize_default_style(doc)

    section = doc.sections[0]
    setup_page(section)
    add_page_number(section)

    api = {
        "h1": add_h1,
        "h2": add_h2,
        "h3": add_h3,
        "para": add_para,
        "code": add_code_block,
        "caption": add_caption,
        "figure_placeholder": add_figure_placeholder,
        "table": add_table,
    }

    front.write(doc, api)
    ch1_intro.write(doc, api)
    ch2_tech.write(doc, api)
    ch3_req.write(doc, api)
    ch4_design.write(doc, api)
    ch5_impl.write(doc, api)
    ch6_test.write(doc, api)
    ch7_summary.write(doc, api)
    back.write(doc, api)

    doc.save(str(OUTPUT))
    print(f"OK output -> {OUTPUT}")


if __name__ == "__main__":
    main()
